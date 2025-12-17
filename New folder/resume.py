from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Helper function to set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_resume():
    doc = Document()
    
    # Set narrow margins (0.5 inches)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Define Colors
    COLOR_PURPLE = RGBColor(44, 68, 111) # #2C446F
    COLOR_BLUE = RGBColor(44, 68, 111)    # #2C446F
    COLOR_JET = RGBColor(42, 45, 49)      # #2A2D31

    # --- Header ---
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run("Niraj Aher")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = COLOR_PURPLE
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("Software Engineer")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = COLOR_BLUE
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_paragraph = doc.add_paragraph()
    contact_text = "Phone: 9769448933 | Email: niraj.aher@gmail.com | LinkedIn: linkedin.com/in/nirajaher |\nAddress: Mumbai / Pune"
    contact_run = contact_paragraph.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = COLOR_JET
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer

    # --- Helper to add Section Headings ---
    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_PURPLE
        p.paragraph_format.space_after = Pt(6)
        # Add a bottom border to the paragraph for the divider look
        p_element = p._p
        pPr = p_element.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '6F5392') # Purple Border
        pbdr.append(bottom)
        pPr.append(pbdr)

    # --- Summary ---
    add_section_heading("SUMMARY")
    summary_text = ("Software Engineering Specialist with 6+ years of experience in Java, PL/SQL, and frameworks like Spring Boot and "
                    "Hibernate. Skilled in developing scalable solutions, optimizing software performance. Proven ability to enhance "
                    "software quality, reduce incidents, and mentor teams to achieve project goals. Certified in Azure Cloud, SQL, Python, and TensorFlow.")
    p = doc.add_paragraph()
    run = p.add_run(summary_text)
    run.font.color.rgb = COLOR_JET
    run.font.size = Pt(10)
    
    # --- Skills ---
    add_section_heading("SKILLS")
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False 
    table.allow_autofit = False
    
    # Set column widths
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(5.5)

    skills_data = [
        ("Programming Languages", "Core Java, Java 8 / J2EE, SQL, PL/SQL, Python"),
        ("Frameworks", "Spring Boot, Spring, Hibernate"),
        ("Web Services / APIs", "RESTful API, JSON, JWT, OAuth 2.0, Swagger"),
        ("Database", "Oracle, Postgres, RDBMS, relational databases"),
        ("Cloud", "Azure, AWS"),
        ("Tools", "Eclipse, SQL Management Studio, JIRA, TortoiseSVN, SonarQube, Maven, Git"),
        ("Machine Learning", "TensorFlow, PyCharm, Pytorch, NumPy, Keras, Scikit-Learn, artificial intelligence")
    ]

    for category, skill_list in skills_data:
        row_cells = table.add_row().cells
        
        # Category (Blue, Bold)
        cat_p = row_cells[0].paragraphs[0]
        cat_run = cat_p.add_run(category)
        cat_run.bold = True
        cat_run.font.color.rgb = COLOR_BLUE
        cat_run.font.size = Pt(10)
        
        # Skill List (Jet)
        skill_p = row_cells[1].paragraphs[0]
        skill_run = skill_p.add_run(skill_list)
        skill_run.font.color.rgb = COLOR_JET
        skill_run.font.size = Pt(10)
    
    doc.add_paragraph() # Spacer

    # --- Work Experience ---
    add_section_heading("WORK EXPERIENCE")

    def add_job(title, company, dates, bullets):
        # Title Line
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        
        title_run = p.add_run(title)
        title_run.bold = True
        title_run.font.color.rgb = COLOR_BLUE
        title_run.font.size = Pt(11)
        
        p.add_run(" - ")
        
        comp_run = p.add_run(company)
        comp_run.bold = True
        comp_run.font.color.rgb = COLOR_JET # Keep company standard or Jet
        comp_run.font.size = Pt(11)
        
        # Tab for date or just new line depending on preference. Using a simpler approach:
        # We will put date on the next line or same line. Let's put date on the right using tabs or just text.
        # Given simpler python-docx control, let's put date on the next line italicized.
        
        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_after = Pt(2)
        date_run = date_p.add_run(dates)
        date_run.italic = True
        date_run.font.color.rgb = COLOR_JET
        date_run.font.size = Pt(9)

        # Bullets
        for b in bullets:
            bp = doc.add_paragraph(style='List Bullet')
            brun = bp.add_run(b)
            brun.font.color.rgb = COLOR_JET
            brun.font.size = Pt(10)
            bp.paragraph_format.space_after = Pt(2)
        
        doc.add_paragraph() # Spacer between jobs

    # Accenture
    add_job("Software Engineering Specialist", "Accenture", "07/2024 - Present", [
        "Wrote secure and high-quality code using Java, gaining guidance from peers to enhance skills and proficiency.",
        "Applied object-oriented programming principles, SOLID design principles, and design patterns to ensure clean and maintainable code.",
        "Designed, developed, and documented RESTful APIs using Spring Boot, Optimized API performance by implementing caching strategies and refining underlying database queries, which reduced average response times by over 20%",
        "Led code reviews, mentored junior developers, and provided training on software development best practices, resulting in a 15% increase in team productivity."
    ])

    # Capgemini
    add_job("Consultant", "Capgemini Technology Services India", "01/2022 - 06/2024", [
        "Accomplished PL/SQL programmer with experience in the development and debugging of packages, procedures, functions, triggers, views, and exception handling in Oracle RDBMS.",
        "Implemented and integrated REST APIs to enhance software functionality and facilitate seamless communication between components.",
        "Led agile backend development teams to ensure the timely delivery of 10+ work orders and meet or exceed all project goals.",
        "Supervised a team of developers and mentored in resolving complex back-end technical issues while monitoring batch performance, resulting in a 20% improvement in issue resolution time.",
        "Collaborated with cross-functional teams to execute basic software solutions through the design, development, and troubleshooting of technical areas.",
        "Analyzed and evaluated software performance, predicted and prioritized potential issues, reducing critical incidents by 30%, and recommended solutions that strengthened the overall quality of the software products.",
        "Optimized batch performance through hands-on experience in performance tuning PL/SQL procedures and functions, leading to a 20% reduction in batch processing time.",
        "Debugged and resolved slow, time-consuming queries, resulting in a 20-25% increase in performance.",
        "Documented software specifications (TDD), drafted code, and reviewed code written by other developers to ensure code quality and consistency, resulting in a 50% decrease in errors and a 30% increase in software performance.",
        "Demonstrated exceptional analytical and problem-solving skills, tackling challenges with a strategic and logical approach."
    ])

    # Deloitte
    add_job("Business Technology Analyst", "Deloitte Consulting, Mumbai", "01/2017 - 10/2019", [
        "Applied automation to reduce manual toil in the Software Development Life Cycle, actively contributing to process efficiency.",
        "Participated in design meetings with peers and stakeholders to ensure effective collaboration and alignment.",
        "Collaborated with the performance testing team and developed utilities using programming languages such as Java, Python, and SQL to automate tasks and reduce performance testing time by around 20 to 30%.",
        "Prioritized tasks effectively, managing time and resources to meet 10 or more work order deadlines while ensuring a high standard of work.",
        "Collaborated in the development and testing of internal REST APIs, writing integration tests with Postman and JUnit to ensure endpoint reliability and data integrity.",
        "Addressed basic and scalable code quality issues, demonstrating attention to detail and commitment to delivering high-quality software.",
        "Applied principles of Object-Oriented Concepts & Data Structures to optimize software design and development.",
        "Successfully fixed 50+ software bugs in multiple applications using Java, JavaScript, and SQL, resulting in a 30% reduction in errors and a 40% improvement in customer satisfaction.",
        "Utilized SVN version control tools to maintain code repositories."
    ])

    # --- Personal Project ---
    add_section_heading("PERSONAL PROJECT")
    
    # Project Title
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("Project Name: Machine Learning model for Text Emotion Recognition")
    run_title.bold = True
    run_title.font.color.rgb = COLOR_BLUE
    run_title.font.size = Pt(11)
    
    # Tech Stack
    p_tech = doc.add_paragraph()
    run_tech = p_tech.add_run("Technologies / Skills: Python, TensorFlow, Pandas, NumPy, matplotlib, Keras, RNN, Data Structure and Algorithms, artificial intelligence")
    run_tech.italic = True
    run_tech.font.color.rgb = COLOR_JET
    run_tech.font.size = Pt(10)
    
    # Description Bullets
    proj_bullets = [
        "The natural language processing (NLP) project utilized deep learning techniques to predict the emotions associated with the analyzed text, resulting in a 95% to 97% accuracy rate.",
        "Trained the model on a large emotion-rich dataset using recurrent neural networks (RNNs) like LSTM for high accuracy.",
        "Conducted data preprocessing, including stop word removal, stemming, and tokenization, resulting in a clean text corpus for 95% accurate analysis and predictions.",
        "GitHub : Text Emotion Recognition"
    ]
    
    for b in proj_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        brun = bp.add_run(b)
        brun.font.color.rgb = COLOR_JET
        brun.font.size = Pt(10)
        bp.paragraph_format.space_after = Pt(2)
        
    doc.add_paragraph() # Spacer

    # --- Certificates ---
    add_section_heading("CERTIFICATES")
    certs = [
        "TensorFlow Developer Certificate (machine learning)",
        "DeepLearning.AI TensorFlow Developer Specialization",
        "Microsoft Certified: Azure Cloud Fundamentals",
        "SQL Certificate",
        "Python Certificate"
    ]
    for c in certs:
        bp = doc.add_paragraph(style='List Bullet')
        brun = bp.add_run(c)
        brun.font.color.rgb = COLOR_JET
        brun.font.size = Pt(10)
        bp.paragraph_format.space_after = Pt(2)
        
    doc.add_paragraph() # Spacer

    # --- Education ---
    add_section_heading("EDUCATION")
    
    # Edu 1
    edu_p1 = doc.add_paragraph(style='List Bullet')
    run_deg1 = edu_p1.add_run("Post Graduate Diploma in Advanced Computing")
    run_deg1.bold = True
    run_deg1.font.color.rgb = COLOR_BLUE
    run_deg1.font.size = Pt(10)
    edu_p1.add_run(" - C-DAC Mumbai").font.color.rgb = COLOR_JET

    # Edu 2
    edu_p2 = doc.add_paragraph(style='List Bullet')
    run_deg2 = edu_p2.add_run("Bachelor of Computer Engineering")
    run_deg2.bold = True
    run_deg2.font.color.rgb = COLOR_BLUE
    run_deg2.font.size = Pt(10)
    edu_p2.add_run(" - Mumbai University").font.color.rgb = COLOR_JET

    doc.save('Niraj_Aher_Resume_Enhanced.docx')

create_resume()